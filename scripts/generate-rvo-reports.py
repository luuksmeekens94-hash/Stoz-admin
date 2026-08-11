from __future__ import annotations

import json
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font

ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "documents" / "concepten" / "report-data-2026-08-11.json"
DOCX_TEMPLATE = ROOT / "documents" / "rvo-templates" / ("Model-D-Voortgangsverslag-STOZ" + ".docx")
XLSX_TEMPLATE = ROOT / "documents" / "rvo-templates" / ("Format-voortgangsverslag-B-STOZ-2025-D2.0" + ".xlsx")
DOCX_OUTPUT = ROOT / "documents" / "concepten" / ("CONCEPT-Model-D-Voortgangsverslag-Hybride-Begrip-2026-08-11" + ".docx")
XLSX_OUTPUT = ROOT / "documents" / "concepten" / ("CONCEPT-Model-B-Financieel-Voortgangsverslag-Hybride-Begrip-2026-08-11" + ".xlsx")

with DATA_PATH.open(encoding="utf-8") as handle:
    DATA = json.load(handle)


def round2(value: float) -> float:
    return round(value + 1e-12, 2)


def nl_money(value: float) -> str:
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def sum_hours(entries: list[dict]) -> float:
    return round2(sum(float(entry["hours"]) for entry in entries))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


def build_model_d() -> None:
    doc = Document(DOCX_TEMPLATE)
    doc.core_properties.title = "Concept voortgangsverslag STOZ – Hybride Begrip"
    doc.core_properties.subject = "RVO Model D – STOZ25-03851282"
    doc.core_properties.author = "Fysiotherapie Fy-fit"
    doc.core_properties.comments = "Concept op basis van de projectadministratie met peildatum 11 augustus 2026."

    doc.paragraphs[6].text = "☐ Ja     ☒ Nee"
    doc.paragraphs[9].text = "☒ Ja     ☐ Nee"
    doc.paragraphs[16].text = "☐ Ja     ☒ Nee"
    doc.paragraphs[20].text = "☐ Ja     ☒ Nee"

    concept_paragraph = doc.paragraphs[2].insert_paragraph_before(
        "CONCEPTVERSIE – peildatum 11 augustus 2026 – nog niet indienen"
    )
    concept_paragraph.style = doc.styles["Normal"]
    for run in concept_paragraph.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    responses = [
        "1 september 2025 tot en met 31 augustus 2026. Dit is een concept op basis van de projectadministratie tot en met 11 augustus 2026. De feitelijke projectuitvoering is in maart 2026 gestart.",
        f"De werkzaamheden zijn niet volledig volgens de oorspronkelijke tijdslijn uitgevoerd. De feitelijke uitvoering startte in maart 2026, later dan de formele projectstart. Sinds maart zijn projectmanagement, inhoudsontwikkeling, digitale productie en de praktijktraining uitgevoerd. De resterende implementatie-, borgings- en evaluatieactiviteiten worden in de operationele forecast vastgelegd met een concrete datum, uitvoerder en uren. De projectadministratie bevat {DATA['forecastEntryCount']} afzonderlijke forecastregels. Op dit moment is geen wijziging van de formele einddatum voorzien.",
        "De latere feitelijke start en het nog beperkte aantal afgeronde cliënttrajecten beïnvloeden de beschikbaarheid van uitkomstgegevens. Daardoor zijn cliëntimpact en medewerkerervaring in deze verslagperiode nog niet kwantitatief te onderbouwen. Daarnaast wijkt de inzet per kostencategorie af van de verleende begroting. Deze afwijkingen worden transparant gerapporteerd; operationeel noodzakelijke inzet wordt niet uit de planning geweerd vanwege een subsidiegrens. Alleen kosten met een herleidbare financiële bron worden als gemaakte kosten opgenomen.",
        "In de verslagperiode zijn geen afzonderlijke externe ontwikkelingen of organisatorische veranderingen geregistreerd die naast de hiervoor beschreven latere feitelijke start aantoonbaar van invloed zijn op de projectuitvoering.",
        "De financiële en inhoudelijke administratie is gereconcilieerd met de ingediende Model B-begroting, de RVO-verleningsbeschikking, urenregistraties, facturen en de presentielijst van de praktijktraining. De beschikking gaat voor waar RVO bedragen heeft aangepast.",
        "Op dit moment wordt in deze conceptversie geen contactverzoek aan een RVO-adviseur opgenomen. Dit kan vóór indiening worden aangepast indien afstemming over een begrotingsafwijking of wijzigingsverzoek wenselijk blijkt.",
        "Er is een digitale meertalige informatievoorziening rond zorgpaden gerealiseerd, met digitale patiëntinformatie en video’s als voorbereiding op en ondersteuning van het behandeltraject. De inhoud en vindbaarheid worden stapsgewijs uitgebreid.",
        "Er is geïnventariseerd, vakinhoud is ontwikkeld en vertaald, digitale content en video zijn geproduceerd en de toepassing is technisch ingericht. De opschaling verloopt gefaseerd per zorgpad, zodat inhoud en werkproces tijdens het gebruik kunnen worden aangescherpt.",
        "De digitale informatie wordt gekoppeld aan intake, behandelplan en vervolgmomenten. Patiënten krijgen gerichte informatie vooraf en tijdens het traject; de fysiotherapeut gebruikt dezelfde content in de begeleiding. Volledige borging in alle reguliere werkprocessen is nog in uitvoering.",
        "Fysiotherapeuten zijn betrokken via een praktijktraining, vakinhoudelijke inhoudsontwikkeling en feedback op toepasbaarheid. De presentielijst en urenadministratie sluiten aan op vijftien aanwezige medewerkers van ieder twee uur. Verdere ondersteuning vindt plaats tijdens de gefaseerde invoering.",
        "Patiënten krijgen eenvoudige meertalige informatie en ondersteunende video’s. De eerste gebruiks- en uitkomstmetingen volgen bij afgeronde trajecten. Op de peildatum zijn nog geen volledige cliënttrajecten met een afgeronde voor- en nameting beschikbaar.",
        "Tot en met 11 augustus 2026 zijn 551,5 goedgekeurde projecturen geregistreerd: 269 uur in WP1, 252,5 uur in WP2 en 30 uur in WP3. Voor WP4, WP5 en WP6 zijn in deze verslagperiode nog geen afzonderlijke goedgekeurde uren geregistreerd.",
        "Tussenresultaten: 15 aanwezige medewerkers hebben de praktijktraining gevolgd (2 uur per deelnemer). Het aantal cliënttrajecten met een volledige voor- en nameting is 0. Er zijn nog geen ingevulde fysiotherapeutenvragenlijsten. Deze aantallen worden in volgende verslagperioden geactualiseerd.",
        "De monitoring combineert cliëntmetingen vóór en na het traject, een meting circa drie maanden na afronding, behandelresultaten en een vragenlijst onder fysiotherapeuten. De registratiebasis is ingericht en wordt gevuld naarmate meer trajecten worden afgerond. De fysiotherapeutenvragenlijst is inhoudelijk voorbereid en wordt verstuurd zodra voldoende implementatie-ervaring is opgedaan.",
        "Niet van toepassing: Hybride Begrip valt onder de opschalingsroute en niet onder de evaluatieroute.",
        "Er zijn nog onvoldoende afgeronde trajecten voor een betrouwbare kwantitatieve impactuitspraak. De verwachte bijdrage ligt in beter begrip, betere voorbereiding, consistente informatie en efficiëntere begeleiding. Dit wordt in volgende verslagperioden getoetst met cliënt- en medewerkergegevens.",
        "Kennisdeling vindt voorlopig binnen Fy-fit plaats via training, gezamenlijke inhoudsontwikkeling en gebruiksfeedback. Externe verspreiding van lessen en herbruikbare werkwijzen is voorzien in WP5 vanaf augustus 2026.",
        "De huidige uitkomsten beschrijven de feitelijke tussenstand en zijn geen eindmeting. Kwantitatieve cliënt- en medewerkeruitkomsten worden pas toegevoegd nadat voldoende volledige trajecten en respons beschikbaar zijn.",
    ]
    if len(doc.tables) < 24:
        raise RuntimeError(f"Onverwacht Model D: {len(doc.tables)} tabellen")
    for index, text in enumerate(responses):
        set_cell_text(doc.tables[index].cell(0, 0), text, size=9)

    activity_names = {activity["code"]: activity["name"] for activity in DATA["activities"]}
    activities = [
        ("A1.1", "Projectmanagement en voortgangssturing", "01-09-2025", "31-08-2027", "In uitvoering", "263 uur geregistreerd; projectsturing en dossieropbouw lopen door."),
        ("A1.2", "Kick-off en gezamenlijke inventarisatie", "01-11-2025", "28-02-2026", "Grotendeels uitgevoerd", "6 uur geregistreerd; inventarisatie is inhoudelijk verwerkt in WP2."),
        ("A2.1", "Digitale basis en meertalige informatie", "01-10-2025", "28-02-2026", "In uitvoering", "80 uur geregistreerd; technische en inhoudelijke basis is aanwezig."),
        ("A2.2", "Vakinhoud ontwikkelen en vertalen", "01-11-2025", "31-03-2026", "In uitvoering", "85 uur geregistreerd; vakinhoud en vertalingen worden per zorgpad aangevuld."),
        ("A2.3", "Video en digitale content produceren", "01-03-2026", "31-05-2026", "In uitvoering", "87,5 uur geregistreerd; eerste content is geproduceerd en wordt verder aangescherpt."),
        ("A3.1", "Praktijktraining medewerkers", "01-12-2025", "31-05-2026", "Uitgevoerd", "15 aanwezigen × 2 uur; presentielijst en urenregistratie zijn gereconcilieerd."),
        ("A3.2", "Vervolgondersteuning en instructie", "01-04-2026", "31-10-2026", "In uitvoering / nog niet apart geregistreerd", "Vervolgondersteuning loopt mee met de gefaseerde implementatie."),
        ("A4.1", "Pilot in reguliere werkprocessen", "01-03-2026", "31-08-2026", "In opstart", "Werkproceskoppeling is beschreven; nog geen afzonderlijke goedgekeurde uren of volledige cliëntmetingen."),
        ("A4.2", "Opschaling naar overige locaties", "01-07-2026", "28-02-2027", "Voorbereiding", "Operationele forecast bevat datum, uitvoerder en uren voor de vervolgfase."),
        ("A5.1", "Kennisdeling en regionale samenwerking", "01-08-2026", "28-02-2027", "Startfase", "Interne kennisdeling is gestart; externe verspreiding volgt in deze fase."),
        ("A5.2", "Borging en bredere verspreiding", "01-03-2027", "31-08-2027", "Nog niet gestart", "Activiteit ligt buiten de huidige feitelijke verslagstand."),
        ("A6.1", "Monitoring en eerste evaluatie", "01-03-2026", "31-08-2027", "In voorbereiding", "Meetopzet is bepaald; cliënt- en fysiotherapeutgegevens zijn nog niet beschikbaar."),
        ("A6.2", "Eindevaluatie en overdracht", "01-03-2027", "31-08-2027", "Nog niet gestart", "Activiteit ligt buiten de huidige feitelijke verslagstand."),
    ]

    table = doc.tables[18]
    while len(table.rows) < len(activities) + 1:
        new_row = table.add_row()
        template_row = table.rows[-2]
        for cell_index, cell in enumerate(new_row.cells):
            cell._tc.get_or_add_tcPr().append(deepcopy(template_row.cells[cell_index]._tc.get_or_add_tcPr()))
    set_repeat_table_header(table.rows[0])
    for row_index, (code, fallback, start, end, status, detail) in enumerate(activities, start=1):
        name = activity_names.get(code, fallback)
        values = [f"{code} · {name}", detail, start, end, status]
        for cell, value in zip(table.rows[row_index].cells, values):
            set_cell_text(cell, value, size=8)

    collaboration_responses = [
        "De samenwerking met de inkoper is in deze verslagperiode nog beperkt. Er zijn nog geen afzonderlijke contractafspraken over de digitale of hybride werkwijze vastgelegd. Borging met inkopers is voorzien binnen WP5 en wordt in de volgende fase concreter uitgewerkt.",
        "Er zijn geen andere formeel deelnemende zorgaanbieders in de financiële projectadministratie opgenomen. Binnen Fy-fit wordt de beoogde inzet geleverd door praktijkhouders, praktijkmanager, fysiotherapeuten, front- en backoffice en de externe projectmanager.",
        "Leveranciers ondersteunen de technische realisatie, website en digitale content. Leveranciersinzet wordt operationeel als uren zichtbaar gehouden, maar wordt in het financiële verslag alleen als gemaakte kosten opgenomen wanneer een factuur of betaalbewijs aanwezig en gekoppeld is.",
        "De samenwerking met overige betrokken partijen staat in deze fase vooral in het teken van interne implementatie. Regionale samenwerking en bredere kennisdeling worden vanaf WP5 verder opgebouwd.",
        "Afspraken met inkopers, externe partners en leveranciers worden in volgende verslagperioden aangevuld zodra deze in de projectadministratie zijn vastgelegd.",
    ]
    for index, text in enumerate(collaboration_responses, start=19):
        set_cell_text(doc.tables[index].cell(0, 0), text, size=9)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "CONCEPT – Hybride Begrip – STOZ25-03851282 – peildatum 11 augustus 2026"
        footer.alignment = 1
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.name = "Arial"

    DOCX_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUTPUT)


def build_model_b() -> None:
    workbook = load_workbook(XLSX_TEMPLATE, data_only=False)
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.calculation.calcMode = "auto"
    cover = workbook["Voorblad"]
    cover["D5"] = "CONCEPT – nog niet indienen"
    cover["D5"].font = Font(name="Arial", size=12, bold=True, color="C00000")
    cover["D6"] = "Peildatum 11 augustus 2026"
    cover["D6"].font = Font(name="Arial", size=10, color="C00000")
    sheet = workbook["Aanvrager-Penvoerder"]

    sheet["C2"] = DATA["applicant"]
    sheet["C3"] = DATA["projectName"]
    sheet["F5"] = "Nee"
    sheet["F6"] = "KMO"
    sheet["F7"] = "N.v.t.; intern begrotingstarief € 50 per uur"
    sheet["F8"] = "Opschalingsroute"
    sheet["F9"] = DATA["approvedSubsidy"]

    entries = DATA["entries"]
    internal_names = set(DATA["reportConfig"]["internalCostUsers"])
    project_entries = [entry for entry in entries if entry["user"] in internal_names and entry["workPackage"] == "WP1"]
    implementation_entries = [entry for entry in entries if entry["user"] in internal_names and entry["workPackage"] not in {"WP1", "WP3"}]

    project_by_person: dict[str, list[dict]] = defaultdict(list)
    implementation_by_person: dict[str, list[dict]] = defaultdict(list)
    for entry in project_entries:
        project_by_person[entry["user"]].append(entry)
    for entry in implementation_entries:
        implementation_by_person[entry["user"]].append(entry)

    input_font = Font(name="Arial", size=10, color="0000FF")

    for row, name in zip(range(16, 25), sorted(project_by_person)):
        sheet[f"B{row}"] = name
        sheet[f"C{row}"] = "Loondienst"
        sheet[f"D{row}"] = 50
        sheet[f"E{row}"] = sum_hours(project_by_person[name])
        for coordinate in (f"B{row}", f"C{row}", f"D{row}", f"E{row}"):
            sheet[coordinate].font = input_font

    for row, name in zip(range(52, 61), sorted(implementation_by_person)):
        sheet[f"B{row}"] = name
        sheet[f"C{row}"] = "Loondienst"
        sheet[f"D{row}"] = 50
        sheet[f"E{row}"] = sum_hours(implementation_by_person[name])
        for coordinate in (f"B{row}", f"C{row}", f"D{row}", f"E{row}"):
            sheet[coordinate].font = input_font

    confirmed_invoices = [
        invoice for invoice in DATA["invoices"]
        if invoice["confirmedBudgetLineId"] == "external-project-manager"
        and invoice["vatTreatment"] == "INCLUDED_CONFIRMED"
        and invoice["hasEvidence"]
    ]
    invoice_ex_vat = round2(sum(invoice["amountExVat"] for invoice in confirmed_invoices))
    invoice_vat = round2(sum(invoice["vatAmount"] for invoice in confirmed_invoices))
    sheet["B68"] = "LS Project- en innovatiemanagement · facturen 66, 67 en 71"
    sheet["D68"] = 100
    sheet["E68"] = round2(invoice_ex_vat / 100)
    sheet["B75"] = "Niet-verrekenbare btw op bevestigde facturen 66, 67 en 71"
    sheet["F75"] = invoice_vat

    present_count = sum(len(training["presentAttendees"]) for training in DATA["trainings"])
    training_hours = DATA["trainings"][0]["hours"] if DATA["trainings"] else 0
    sheet["B102"] = "Praktijktraining Hybride Begrip · aanwezige medewerkers"
    sheet["C102"] = present_count
    sheet["D102"] = 50
    sheet["E102"] = training_hours

    # Verleende bedragen uit ingediende Model B, gecorrigeerd conform RVO-beschikking.
    sheet["I16"] = 16250
    sheet["I52"] = 7250
    sheet["I53"] = 3000
    sheet["I54"] = 1000
    sheet["I68"] = 32500
    sheet["I69"] = 2500
    sheet["I87"] = 8000
    sheet["I88"] = 2400
    sheet["I116"] = 1000
    sheet["I139"] = 645

    note_lines = [
        "CONCEPT – peildatum 11 augustus 2026. Nog niet indienen.",
        "Verleende basis: ingediende Model B-begroting en RVO-beschikking STOZ25-03851282; de beschikking gaat voor bij aangepaste bedragen.",
        "Fy-fit kan btw voor dit project niet verrekenen. Niet-verrekenbare btw op bevestigde facturen is daarom als projectkosten opgenomen.",
        f"Externe projectmanagementkosten zijn gebaseerd op facturen 66, 67 en 71: € {nl_money(invoice_ex_vat)} exclusief btw en € {nl_money(invoice_vat)} btw.",
        "Website-inzet en overige leveranciersinzet zonder gekoppelde factuur of betaalbewijs zijn niet als gemaakte financiële kosten opgenomen, ook als operationele uren zichtbaar zijn.",
        "Fysiotherapie-uren voor inhoudsontwikkeling zijn onder Implementatie opgenomen; vijftien aanwezige cursisten van twee uur zijn onder Opleiding opgenomen.",
        "Begrotingsafwijkingen vragen vóór definitieve indiening nog een bestuurlijke controle en zo nodig afstemming met RVO.",
    ]
    for row, line in enumerate(note_lines, start=161):
        cell = sheet.cell(row=row, column=2)
        cell.value = line
        cell.number_format = "@"
        cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
        cell.font = Font(name="Arial", size=9, color="0000FF")
        sheet.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        sheet.row_dimensions[row].height = 30 if len(line) < 140 else 42

    for worksheet in workbook.worksheets:
        worksheet.sheet_view.showGridLines = False
    for index in range(1, 8):
        participant = workbook[f"Deelnemer{index}"]
        participant.sheet_state = "hidden"
        for coordinate in ("J43", "J77", "J91", "J138", "J140"):
            formula = participant[coordinate].value
            if isinstance(formula, str) and formula.startswith("=") and not formula.startswith("=IFERROR("):
                participant[coordinate] = f"=IFERROR({formula[1:]},0)"
    workbook.save(XLSX_OUTPUT)


if __name__ == "__main__":
    build_model_d()
    build_model_b()
    print(json.dumps({"docx": str(DOCX_OUTPUT), "xlsx": str(XLSX_OUTPUT)}, indent=2))
